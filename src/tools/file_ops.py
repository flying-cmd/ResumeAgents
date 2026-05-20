import os
from langchain_core.tools import tool
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from rich.console import Console

console = Console()


def convert_doc_to_docx(doc_path: str) -> str:
    """
    Convert a legacy .doc file to .docx with win32com on Windows.
    Returns the path to the converted .docx file.
    """
    try:
        import win32com.client as win32
        import pythoncom
    except ImportError:
        raise ImportError("pywin32 is not installed. Please install it with `pip install pywin32`.")

    pythoncom.CoInitialize()

    word = win32.Dispatch("Word.Application")
    word.Visible = False

    doc_abspath = os.path.abspath(doc_path)
    docx_path = doc_abspath + "x"

    try:
        wb = word.Documents.Open(doc_abspath)
        wb.SaveAs2(docx_path, FileFormat=12)
        wb.Close()
        word.Quit()
        return docx_path
    except Exception as e:
        try:
            word.Quit()
        except Exception:
            pass
        raise e


@tool
def read_tech_stack_tool(file_path: str) -> str:
    """Read the technology stack from a text file."""
    if not os.path.exists(file_path):
        return f"Error: Technology stack file not found at {file_path}"

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        return content
    except Exception as e:
        return f"Error reading technology stack: {str(e)}"


@tool
def read_resume_tool(directory: str) -> str:
    """
    Read the resume from a PDF, DOCX, or DOC file in the target directory.
    The file extension is detected automatically.
    """
    if not os.path.exists(directory):
        return f"Error: Directory not found at {directory}"

    files = os.listdir(directory)
    resume_files = [f for f in files if f.lower().endswith((".pdf", ".docx"))]
    legacy_files = [f for f in files if f.lower().endswith(".doc")]
    resume_files.sort(key=lambda name: os.path.getmtime(os.path.join(directory, name)), reverse=True)
    legacy_files.sort(key=lambda name: os.path.getmtime(os.path.join(directory, name)), reverse=True)

    file_to_process = None

    if resume_files:
        file_to_process = os.path.join(directory, resume_files[0])
    elif legacy_files:
        file_to_process = os.path.join(directory, legacy_files[0])
    else:
        return "Error: No PDF or DOCX resume found in the data directory."

    file_ext = os.path.splitext(file_to_process)[1].lower()

    try:
        content = ""
        if file_ext == ".pdf":
            loader = PyPDFLoader(file_to_process)
            pages = loader.load()
            content = "\n".join([page.page_content for page in pages])
        elif file_ext == ".docx":
            loader = Docx2txtLoader(file_to_process)
            docs = loader.load()
            content = "\n".join([doc.page_content for doc in docs])
        elif file_ext == ".doc":
            console.print("[yellow][Extractor Agent] Legacy .doc file detected. Attempting automatic conversion to .docx...[/yellow]")
            try:
                temp_docx = convert_doc_to_docx(file_to_process)
                loader = Docx2txtLoader(temp_docx)
                docs = loader.load()
                content = "\n".join([doc.page_content for doc in docs])

                try:
                    os.remove(temp_docx)
                except Exception:
                    pass

                console.print("[green][Extractor Agent] .doc conversion succeeded.[/green]")
            except ImportError:
                return "Error: Found a .doc file but pywin32 is not installed, so the legacy format cannot be converted."
            except Exception as e:
                return (
                    "Error: Failed to convert the .doc file. Please verify that Microsoft Word is installed "
                    f"or convert the file manually. Details: {str(e)}"
                )
        else:
            return f"Error: Unsupported file format {file_ext}"

        return content
    except Exception as e:
        return f"Error reading resume: {str(e)}"


@tool
def save_document_tool(content: str, file_path: str) -> str:
    """Save text content to a file."""
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Successfully saved document to {file_path}"
    except Exception as e:
        return f"Error saving document: {str(e)}"


@tool
def generate_docx_tool(markdown_content: str, output_path: str) -> str:
    """Convert Markdown content to a Word (.docx) file with improved formatting."""
    console.print(f"[cyan][Word Generator][/cyan] Starting DOCX generation: {os.path.basename(output_path)}...")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        lines = markdown_content.split("\n")

        def add_formatted_run(paragraph, text):
            """Apply simple bold formatting for **text** segments."""
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                heading = doc.add_heading(level=1)
                run = heading.add_run(line[2:])
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(44, 62, 80)
                run.bold = True
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            elif line.startswith("## "):
                heading = doc.add_heading(level=2)
                run = heading.add_run(line[3:])
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(52, 73, 94)
                run.bold = True

            elif line.startswith("### "):
                heading = doc.add_heading(level=3)
                run = heading.add_run(line[4:])
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(127, 140, 141)
                run.bold = True

            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_run(p, line[2:])

            else:
                p = doc.add_paragraph()
                add_formatted_run(p, line)

        doc.save(output_path)
        msg = f"Successfully generated DOCX at {output_path}"
        console.print("[green][Word Generator][/green] DOCX generation completed.")
        return msg

    except ImportError:
        msg = "Error: python-docx is not installed. Please install it with `pip install python-docx`."
        console.print(f"[bold red]{msg}[/bold red]")
        return msg
    except Exception as e:
        msg = f"Exception during DOCX generation: {str(e)}"
        console.print(f"[bold red]{msg}[/bold red]")
        return msg
